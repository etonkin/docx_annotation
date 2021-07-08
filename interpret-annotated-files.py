#!/opt/python/anaconda-3.7/bin/python
# -*- coding: utf-8 -*-
from __future__ import generators
import docx
import os
import sqlite3
from sqlite3 import Error
import sys;
import nltk;
from nltk.tokenize import TweetTokenizer
from os.path import isfile, join
import json
from nltk import word_tokenize
import re
from os import listdir
import getopt
import logging


def print_help():
    print("Help:")
    print("--help or -h: print this information");
    print("--input=docxfile.docx or -i docxfile.docx: give file on which to operate)");
    print("--dirinput=dirname or -d dirname: scan a whole directory of docx files")
    print("--verbose or -v: tell me everything about what you are doing")
    print("--scan or -s: look for and store highlights")
    print("--colourschema or -c: give file containing equivalence schema between annotations and colours (e.g. RED,ANGRY)")
    print("--output or -o: directory into which to store extracted JSONs")
    print("--database=database.db: database to lookup text location in database/document store (will not be calculated if empty)");
    print("--updatedb or -u: update the database to contain relevant annotations (will not be done if not set")

def encode_highlight_colors(highlight_color,schema):
    highlight_color=str(highlight_color);
    for c in schema.keys():
        if highlight_color.startswith(c):
            return schema[c];
    return 'UNSET';

def nuke_fancy_quotes(text):
    # unfortunately windows uses an idiosyncratic codepage so fills documents with fancy quotes (see https://en.wikipedia.org/wiki/Windows-1252#Codepage_layout)
    # for those relying on comparison with tokenised documents, this can cause complications. So we try to replace these with something more standard. 
    text=text.replace(u'\u201c', '"').replace(u'\u201d', '"') 
    text=text.replace(u'\u2018', "'").replace(u'\u2019', "'") 
    text=text.replace(u'\u2013', "-").replace(u'\u2014', "-") 
    text=text.replace(u'\u20ac', "€")
    text=text.replace(u'\u00a3', "£ ") 
    # idiosyncratic changes to render the behaviour of tokenizer closer to that of ucto
    # where you have ' before an acronym, add a space
    re_outer = re.compile(r'([^A-Z ])([A-Z])')
    re_inner = re.compile(r'\b[A-Z]+(?=[A-Z][a-z])')
    text = re.sub( r'([\'])([A-Z][A-Z])', r'\1 \2',text)
    return(text);

def uctolike_tokenize(text):
    tokenized_text=word_tokenize(text);
    # weirdness of this tokenizer: replaces cannot with can not, two words. 
    # for compatibility, doing the same here
    # This is clearly very idiosyncratic/contextual. 
    tokenized_str=' '.join(tokenized_text)
    tokenized_str=tokenized_str.replace("can not","cannot")
    tokenized_str=tokenized_str.replace("First-tier","First - tier")
    tokenized_alt_text=tokenized_str.split();
    return tokenized_text,tokenized_alt_text; 


def process_word_file(docurl,filename,schema):
    #document=docx.Document(docurl);
    tknzr = TweetTokenizer()
    document=docx.Document(docurl);
    #print(docurl);
    highlights = []
    for paragraph in document.paragraphs:
        interpreted_highlight={};
        #print(paragraph);
        highlight = ""
        cur_highlight_col=False;
        # store the paragraph the text came from
        # would be handy to also store just the sentence(s) intersecting with the highlight
        partext=paragraph.text;
        logging.debug("PARAGRAPH TEXT");
        logging.debug(partext);
        
        for run in paragraph.runs:
            logging.debug(run.text); 
            interpreted_highlight['text']="";
            # With a simplisitic approach a bug shows up if two highlighted segments appear together, but with a gap (like a blank space which is not highlighted). 
            # runs include all text spans, both highlighted and not. 
            if run.font.highlight_color:
                if(run.font.highlight_color!=cur_highlight_col and cur_highlight_col!=False):
                    # highlight colour change has occurred
                    logging.debug("Highlight color change!");
                    logging.debug("Current state of highlight:");
                    logging.debug(highlight);
                    # this happens if highlight colour changes
                    # kill this one off and start a new one. 
                    interpreted_highlight['text']=nuke_fancy_quotes(highlight);
                    # tokenise a copy of it. This will probably come in handy later. 
                    interpreted_highlight['alt_tokenised_text'],interpreted_highlight['tokenised_text']=uctolike_tokenize(interpreted_highlight['text']);
                    logging.debug("ALT TEXT");
                    logging.debug(interpreted_highlight['alt_tokenised_text'])
                    logging.debug("ORIG TEXT");
                    logging.debug(interpreted_highlight['tokenised_text'])
                    interpreted_highlight['tag']=encode_highlight_colors(run.font.highlight_color,schema);
                    interpreted_highlight['textcolor']=run.font.highlight_color;
                    interpreted_highlight['file_origin']=filename 
                    interpreted_highlight['textualcontext']=partext;
                    highlights.append(interpreted_highlight);
                    interpreted_highlight={};
                    highlight="";
                cur_highlight_col=run.font.highlight_color;
                highlight += run.text
                interpreted_highlight['textcolor']=run.font.highlight_color;
                interpreted_highlight['tag']=encode_highlight_colors(run.font.highlight_color,schema);
                interpreted_highlight['file_origin']=filename 
                interpreted_highlight['textualcontext']=partext;
            else:
               if (highlight!=""):
            	    # no longer have highlight, so the span must've finished. Therefore, shove it in a stored highlight
                   logging.debug("Current state of highlight (on false)");
                   logging.debug(highlight); 
                   cur_highlight_col=False;
                   interpreted_highlight['file_origin']=filename 
                   interpreted_highlight['textualcontext']=partext;
                   interpreted_highlight['text']=nuke_fancy_quotes(highlight);
                   interpreted_highlight['alt_tokenised_text'],interpreted_highlight['tokenised_text']=uctolike_tokenize(interpreted_highlight['text']);
                   logging.debug("ALT TEXT");
                   logging.debug(interpreted_highlight['alt_tokenised_text'])
                   logging.debug("ORIG TEXT");
                   logging.debug(interpreted_highlight['tokenised_text'])
                   highlights.append(interpreted_highlight);
                   interpreted_highlight={};
                   highlight="";
        if highlight:
            interpreted_highlight['file_origin']=filename 
            interpreted_highlight['textcolor']=run.font.highlight_color;
            interpreted_highlight['text']=nuke_fancy_quotes(highlight);
            interpreted_highlight['alt_tokenised_text'],interpreted_highlight['tokenised_text']=uctolike_tokenize(interpreted_highlight['text']);
            highlights.append(interpreted_highlight);

    return highlights;

def create_connection(db_file):
    """ create a database connection to a SQLite database """
    conn = None
    try:
        conn = sqlite3.connect(db_file)
        print(sqlite3.version)
    except Error as e:
        print(e)
    finally:
        if conn:
            return conn;

def KnuthMorrisPratt(text, pattern):

    # (PSF-licenced) https://code.activestate.com/recipes/117214-knuth-morris-pratt-string-matching/ 
    '''Yields all starting positions of copies of the pattern in the text.
Calling conventions are similar to string.find, but its arguments can be
lists or iterators, not just strings, it returns all matches, not just
the first one, and it does not need the whole text in memory at once.
Whenever it yields, it will have read the text exactly up to and including
the match that caused the yield.'''

    # allow indexing into pattern and protect against change during yield
    pattern = list(pattern)

    # build table of shift amounts
    shifts = [1] * (len(pattern) + 1)
    shift = 1
    for pos in range(len(pattern)):
        while shift <= pos and pattern[pos] != pattern[pos-shift]:
            shift += shifts[pos-shift]
        shifts[pos+1] = shift

    # do the actual search
    startPos = 0
    matchLen = 0
    for c in text:
        while matchLen == len(pattern) or \
              matchLen >= 0 and pattern[matchLen] != c:
            startPos += shifts[matchLen]
            matchLen -= shifts[matchLen]
        matchLen += 1
        if matchLen == len(pattern):
            yield startPos

def find_sublist(sub, bigger):
    if not bigger:
        return -1
    if not sub:
        return 0
    first, rest = sub[0], sub[1:]
    pos = 0
    try:
        while True:
            pos = bigger.index(first, pos) + 1
            if not rest or bigger[pos:pos+len(rest)] == rest:
                return pos -1 
    except ValueError:
        return -1

def scan_directory(thedirname):
    scan_results=[];
    onlyfiles = [f for f in listdir(thedirname) if isfile(join(thedirname, f)) ]
    for f in onlyfiles:
        lowerf=f.lower();
        if lowerf.endswith('.docx'):
            scan_results.append(f)
    return scan_results;

def read_schema(theschema):
    f = open(theschema, "r")
    schema={};
    for l in f:
        try:
            l=l.rstrip();
            a,b=l.split(',');
            print(a);
            schema[a]=b;
        except:
            pass;
    return(schema);
    
if __name__ == "__main__":
    # Get command line options 
    logging.basicConfig(filename='docx-annotations.log', level=logging.INFO)
    arguments = len(sys.argv) - 1
    full_cmd_arguments = sys.argv
    argument_list = full_cmd_arguments[1:]
    short_opts="hi:d:vsrb:o:c:" 
    long_opts=["help","input=","dirinput=","verbose","output=","outdir=","colourschema="]; 
    inputdir="";
    database="";
    outdir="";
    schemalocation="";
    updatedb=False;
    schema={};
    try:
        arguments, values = getopt.getopt(argument_list, short_opts, long_opts)
    except getopt.error as err:
        # Output error, and return with an error code
        print (str(err))
        sys.exit(2)
    target_files=[];
    for current_argument, current_val in arguments:
        logging.debug(current_argument+", "+current_val)
        if(current_argument=="--verbose" or current_argument=="-v"):
            debuglevel="debug";
            logging.basicConfig(filename='sample.log', level=logging.DEBUG)
        if(current_argument=="--help" or current_argument=="-h"):
            print_help();
            sys.exit();

    for current_argument, current_val in arguments:
        if(current_argument=="--colourschema" or current_argument=="-c"):
            # this should be comma separated variables, COLOUR,APPLICABLE_TAG
            schema=read_schema(current_val);
            print(schema);
        if(current_argument=="--output" or current_argument=="-o" or current_argument=="--outdir"):
            outdir=current_val;
            print("Extracting files to "+outdir);
            if(not os.path.isdir(outdir)):
                print("can't find directory "+outdir);
                sys.exit();
        if(current_argument=="--updatedb" or current_argument=="-u"):
            updatedb=True;
        if(current_argument=="--database" or current_argument=="-b"):
            database=current_val;
            logging.debug("Will try to use database "+database);
            conn=create_connection(current_val)
            curr=conn.cursor();
        if(current_argument=="--input" or current_argument=="-i"):
            target_files.append(current_val);
            logging.debug(str(target_files));
        if(current_argument=="--dirinput" or current_argument=="-d"):
            # scan through directory current_val to look for docx files, add all to list
            if(not os.path.isdir(current_val)):
                print("Object to scan must be directory")
                sys.exit()
            else:
                inputdir=current_val;
                print("Scanning")
                target_files=scan_directory(current_val);

    if(len(schema.keys())==0):
        print("No schema available. Just using colours.");
    sql_seek_file_prefix= """ SELECT * FROM cases where filename=? """
    sql_get_words_from_file=""" select * from corpus where fileid=? """
    sql_update_tags=""" UPDATE corpus set tag=? where id=?"""
    for filename in target_files:
        file_to_scan=os.path.join(inputdir,filename);
        highlightset=[];
        highlights=process_word_file(os.path.join(inputdir,filename),filename,schema);
        annotationcount=0;
        for i in highlights:
            highlightset.append(i);
        for i in highlightset:
            tag=i['tag'];
            logging.debug(i);
            tokenised_terms=i['tokenised_text'];
            alt_tokenised_terms=i['alt_tokenised_text'];
            logging.debug(i);
            if len(database) > 0 :
                # solely local convention of no relevance to most users 
                f=i['file_origin'].strip('.docx')+'.txt.ucto';
                curr.execute(sql_seek_file_prefix,(f,));    
                rows=curr.fetchall();
                itemid="";
                for row in rows:
                    itemid=row[0];
                if(itemid!=""):
                    # we know which item it is
                    # so now we need to find the corresponding segment from all rows 
                    curr.execute(sql_get_words_from_file,(itemid,));    
                    rows=curr.fetchall();
                    wordset=[];
                    wordid=[];
                    fileid=[];
                    sentenceid=[];
                    existing_tagset=[];
                    for row in rows:
                        wordid.append(row[0]);
                        fileid.append(row[1]);
                        sentenceid.append(row[2]);
                        word=row[3];
                        wordset.append(word); 
                        existing_tagset.append(row[5])
                        # this often fails because of microsoft word fancy quotes, so we remove these.
                        # (this has been implemented here, see nuke fancy quotes method - there are probably other things that could trip this under certain circs so it is possible some cases will fail)
                        # choosing one of these methods entirely at random (should time them one day)
                        #for s in KnuthMorrisPratt(wordset,tokenised_terms): 
                        index= find_sublist(tokenised_terms,wordset)
                        if(index==-1):
                            # second try with what we know are a couple of common issues (these tend to relate to differences in tokenisation practices which mess with the DB)
                            index=find_sublist(alt_tokenised_terms,wordset);
                        tokenisedterm_len=len(tokenised_terms);
                        if(index!=-1): 
                            i['doc-word-index']=index;
                            i['corpus_start_wordid']=wordid[index];
                            i['corpus_start_sentenceid']=sentenceid[index];
                            i['corpus_fileid']=fileid[index];
                            if(updatedb): 
                                for x in range(index,index+tokenisedterm_len):
                                    curr.execute(sql_update_tags,(tag,wordid[x],));
                                    conn.commit();
                        else: # couldn't work out the index , just record it without any position in file information - the -1 in the filename means 'got stuck here'.
                            i['corpus_start_wordid']=-1;
                            i['doc-word-index']=-1;
           # if no db, we haven't been able to work out the position in the ucto-tokenised dataset because we didn't have a db to look it up in - 
	   # so we just give the information we have without it 
            if( 'corpus_start_wordid' in i.keys()): #if we know where this actually came from in the file
                fname="msword-decision-"+i['file_origin']+'-'+str(i['corpus_start_wordid'])+'.json'; 
            else: # no idea where it comes from in the file, just name it with an incrementing counter
                fname="msword-decision-"+i['file_origin']+'-no-'+str(annotationcount)+'.json';
                annotationcount=annotationcount+1;
            with open (os.path.join(outdir,fname), 'w') as outfile:
                json.dump(i,outfile);
        annotationcount=0;
#vim: ts=4 sw=4 et
